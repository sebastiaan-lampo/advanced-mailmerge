import re

import pandas as pd
import docx
import logging
from docx.shared import RGBColor, Cm
from util_functions import add_bookmark, add_bookmark_pageref, set_cell_color, add_bookmark_ref

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger('playbook')


def apply_style(cell, style):
    styles = {
        # Color, Text Color, Bold?
        'Header': ('2F5496', 'FFFFFF', True),
        'Subheader': ('4B84E8', 'FFFFFF', False),
        'Subsubheader': ('BDBDBD', '000000', False)
    }
    assert style in styles
    c, tc, bold = styles[style]
    set_cell_color(cell, c)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = bold
            r.font.color.rgb = RGBColor.from_string(tc)


def load_data(filename, sheet=None):
    with open(filename, 'rb') as f:
        df = pd.read_excel(f, sheet_name=sheet)
    return df


def add_acronyms(doc, df, lookup):
    acronyms = {'Sr AM', 'AM TPM', 'AIC', 'ARQTS'}
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            content = str(df.iloc[i, j])
            if content == "nan":
                continue

            res = re.findall('[A-Z][a-z]*[A-Z]+[a-z]*', content)
            if res:
                for a in res:
                    acronyms.add(a)
    acronyms = sorted(acronyms)
    for a in acronyms:
        try:
            acronyms.remove(f'{a}s')  # remove plurals
        except ValueError:
            pass
    logger.debug(f'Acronyms: {acronyms}')

    doc.add_heading('Acronyms')
    for a in acronyms:
        p = doc.add_paragraph('', style='List Bullet')
        r = p.add_run(a)
        r.font.bold = True
        try:
            d = lookup.loc[a, "Definition"]
        except KeyError:
            logger.debug(f'Missing acronym: {a}')
            d = ""
        p.add_run(f':\t{d}')
    doc.add_page_break()
    return doc


def add_defined_terms(doc, df):
    defined_terms = set()
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            content = str(df.iloc[i, j])
            if content == "nan":
                continue

            res = re.findall('([A-Z]\w+\W([A-Z]\w+\W)+)', content)
            if res:
                for a in res:
                    defined_terms.add(a[0])
    defined_terms = sorted(defined_terms)
    logger.debug(f'Defined terms: {defined_terms}')

    doc.add_heading('Capitalized sequences')
    for a in defined_terms:
        p = doc.add_paragraph('')
        r = p.add_run(a)
        # r.font.bold = True
        # try:
        #     d = lookup.loc[a, "Definition"]
        # except KeyError:
        #     logger.debug(f'Missing acronym: {a}')
        #     d = ""
        # p.add_run(f':\t{d}')
    doc.add_page_break()
    return doc


def add_tasks(doc, df):
    from docx.shared import Cm
    logger.debug(f'labels in task dataframe: {df.columns.values}')
    doc.add_heading('Tasks', level=1)
    df = df.reset_index()

    def add_single_task_layout(tbl, df_row):
        rows = [tbl.add_row() for _ in range(4)]
        for r in rows:
            r.AllowBreakAcrossPages = False

        rows[0].cells[0].merge(rows[3].cells[2])
        p = rows[0].cells[0].paragraphs[0]
        add_bookmark(p, str(df_row["Reference #"]), str(df_row["Reference #"]))
        p.add_run(f'. {df_row["Task (label in flowchart)"]}')
        # rows[0].cells[0].text = f'{df_row["Reference #"]}. {df_row["Task (label in flowchart)"]}'
        for r in p.runs:
            r.font.bold = True
        # rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        rows[0].cells[0].add_paragraph(df_row["Considerations"])
        if df_row["Contract model considerations"] != "Independent of delivery model.":
            p = rows[0].cells[0].add_paragraph(f'Contract model considerations')
            p.add_run(f': {df_row["Contract model considerations"]}')
            p.runs[0].font.bold = True
        for p in rows[0].cells[0].paragraphs:
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True

        for offset, col in enumerate(['Sr AM', 'AM TPM', 'AIC', 'ARQTS']):
            rows[offset].cells[3].text = col
            rows[offset].cells[4].text = df_row.iloc[len(df_row.index.values) - 4 + offset]
            if rows[offset].cells[4].text == "nan":
                rows[offset].cells[4].text = ""
        return tbl

    different = df.ne(df.shift())
    # logger.debug(different.head())
    i = 0
    while i < df.shape[0]:
        if different.loc[i, "Phase"]:
            doc.add_paragraph(df.loc[i, "Phase"], style="Heading 2")

        tbl = doc.add_table(rows=3, cols=0)
        widths = [2, 5, 5, 2, 1]
        for w in widths:
            tbl.add_column(width=Cm(w))

        tbl.style = 'Table Grid'  # 'Light Shading Accent 1'
        tbl.cell(0, 0).text = "Goal"
        tbl.cell(0, 1).merge(tbl.cell(0, 4))
        tbl.cell(0, 1).text = df.loc[i, "Goal / Risk"]
        tbl.cell(1, 0).text = "Objective"
        tbl.cell(1, 1).merge(tbl.cell(1, 4))
        tbl.cell(1, 1).text = df.loc[i, "Objective"]
        tbl.cell(2, 0).merge(tbl.cell(2, 4))
        tbl.cell(2, 0).text = "Tasks and RACI"

        apply_style(tbl.cell(0, 0), 'Header')
        apply_style(tbl.cell(0, 1), 'Header')
        apply_style(tbl.cell(1, 0), 'Subheader')
        apply_style(tbl.cell(1, 1), 'Subheader')
        apply_style(tbl.cell(2, 0), 'Subsubheader')

        tbl = add_single_task_layout(tbl, df.iloc[i, :])

        n = 1
        try:
            while not different.loc[i + n, "Objective"]:
                tbl = add_single_task_layout(tbl, df.iloc[i + n, :])
                n += 1
        except KeyError:  # End of the table.
            pass
        i += n

        # Word seems to ignore this property for merged cells, unfortunately.
        for r in tbl.rows:
            r.AllowBreakAcrossPages = False

        doc.add_paragraph()

    doc.add_page_break()
    return doc


def add_theme_breakdown(doc, df):
    logger = logging.getLogger('playbook.theme')
    logger.setLevel(logging.DEBUG)
    doc.add_heading('Tasks by theme', level=1)
    themes = df['Theme'].unique()
    df = df.reset_index()

    def add_task_reference(tbl, df, index):
        r = tbl.add_row()
        r.cells[2].text = df.loc[index, "Task (label in flowchart)"]
        add_bookmark_ref(r.cells[0].paragraphs[0], str(df.loc[index, "Reference #"]))
        add_bookmark_pageref(r.cells[1].paragraphs[0], str(df.loc[index, "Reference #"]))
        r.cells[3].text = df.loc[index, "Senior AM"]
        r.cells[4].text = df.loc[index, "AM Tech PM"]
        r.cells[5].text = df.loc[index, "Asset Information Coordinator"]
        r.cells[6].text = df.loc[index, "ARQTS"]

    columns = [
        ('Ref #', 1.5),
        ('Page', 1.5),
        ('Task', 9),
        ('Sr AM', 1.5),
        ('AM TPM', 1.5),
        ('AIC', 1.5),
        ('ARQTS', 1.5)
    ]
    for theme in themes:
        df_theme = df[df['Theme'] == theme]
        doc.add_heading(theme, level=2)
        tbl = smartly_add_table(doc, columns)

        for i in df_theme.index.values:
            add_task_reference(tbl, df_theme, i)
        doc.add_page_break()

    return doc


def add_role_breakdown(doc, df):
    logger = logging.getLogger('playbook.role')
    logger.setLevel(logging.DEBUG)
    doc.add_heading('Tasks by role', level=1)
    roles = ['Senior AM', 'AM Tech PM', 'Asset Information Coordinator', 'ARQTS']

    df = df.reset_index()

    def add_task_reference(tbl, df, index, role):
        r = tbl.add_row()
        r.cells[2].text = df.loc[index, "Task (label in flowchart)"]
        add_bookmark_ref(r.cells[0].paragraphs[0], str(df.loc[index, "Reference #"]))
        add_bookmark_pageref(r.cells[1].paragraphs[0], str(df.loc[index, "Reference #"]))
        r.cells[3].text = df.loc[index, role]

    for role in roles:
        columns = [
            ('Ref #', 1.5),
            ('Page', 1.5),
            ('Task', 9.5),
            (role, 3),
        ]

        df_role = df[df[role] != '']
        doc.add_heading(role, level=2)
        tbl = smartly_add_table(doc, columns)

        for i in df_role.index.values:
            add_task_reference(tbl, df_role, i, role)
        doc.add_page_break()

    return doc


def smartly_add_table(doc, columns):
    """
    Add a table with headers and defined column width. Apply document style too.
    :param doc:
    :param columns: List of tubples with ('column name', int) with int == width in cm.
    :return: newly created table reference
    """

    tbl = doc.add_table(rows=1, cols=0)
    for _, width in columns:
        tbl.add_column(width=Cm(width))
    tbl.style = 'Table Grid'  # 'Light Shading Accent 1'
    for col_def, cell in zip(columns, tbl.rows[0].cells):
        label, _ = col_def
        set_cell_color(cell, '2F5496')
        cell.text = label
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    return tbl


def add_phase_breakdown(doc, df):
    # from docx.shared import Cm
    doc.add_heading('Tasks by phase', level=1)

    df = df.reset_index()
    different = df.ne(df.shift())
    i = 0
    while i < df.shape[0]:
        if different.loc[i, "Phase"]:
            doc.add_paragraph(df.loc[i, "Phase"], style="Heading 2")

        columns = [
            ('Ref #', 1.5),
            ('Page', 1.5),
            ('Task', 9),
            ('Sr AM', 1.5),
            ('AM TPM', 1.5),
            ('AIC', 1.5),
            ('ARQTS', 1.5)
        ]
        tbl = smartly_add_table(doc, columns)

        def add_task_reference(index):
            r = tbl.add_row()
            r.cells[2].text = df.loc[index, "Task (label in flowchart)"]
            add_bookmark_ref(r.cells[0].paragraphs[0], str(df.loc[index, "Reference #"]))
            add_bookmark_pageref(r.cells[1].paragraphs[0], str(df.loc[index, "Reference #"]))
            r.cells[3].text = df.loc[index, "Senior AM"]
            r.cells[4].text = df.loc[index, "AM Tech PM"]
            r.cells[5].text = df.loc[index, "Asset Information Coordinator"]
            r.cells[6].text = df.loc[index, "ARQTS"]

        add_task_reference(i)
        n = 1
        try:
            while not different.loc[i + n, "Phase"]:
                add_task_reference(i + n)
                n += 1
        except KeyError:  # End of the table.
            pass
        i += n
        doc.add_page_break()

    return doc


def new_info_only_sheet(filename, df):
    # from openpyxl import load_workbook

    # df = df.set_index(keys='Reference #')
    different = df[df.ne(df.shift())]
    for c in ['Contract model considerations', 'Senior AM', 'AM Tech PM', 'Asset Information Coordinator', 'ARQTS']:
        different[c] = df.loc[:, c]

    # Last 4/5 columns should be carried entirely, they are the RACI. Adjust number as needed.
    # NOT RELIABLE. Columns could be ordered differently.
    # for c in range(5):
    #     different.iloc[:, different.shape[1] - c-1] = df.iloc[:, different.shape[1] - c-1]

    logger.info(different.columns.values)
    with pd.ExcelWriter(filename, mode='w') as writer:
        different.to_excel(writer, sheet_name='detail_grouped')


def extract_comments(filename, sheetname='WKT'):
    from openpyxl import load_workbook
    wb = load_workbook(filename)
    ws = wb[sheetname]

    comments = []
    for _, cell in ws._cells.items():
        if cell.comment:
            logger.debug((cell.row, cell.column, ws.cell(row=1, column=cell.column).value, cell.comment.text))
            comments.append((cell.row, cell.column, ws.cell(row=1, column=cell.column).value, cell.comment.text))

    return comments


def add_comments(doc, df, comments):
    doc.add_heading('WKT Comments', level=2)
    for comment in comments:
        row, col, col_label, text = comment
        row = row-1
        p = doc.add_paragraph('', style='List Bullet')
        r = p.add_run(f'Cell {chr(ord("@")+col)}{row+1} ({col_label}):')
        r.font.bold = True
        r = p.add_run(text)
        r.add_break()
        if row >= 1:
            p.add_run(f'Cell contents: {str(df.loc[row, col_label])}')

    doc.add_page_break()
    return doc


if __name__ == '__main__':
    df_wkt = load_data('playbook_wkt.xlsx', 'WKT')
    acronyms = load_data('playbook_wkt.xlsx', 'Acronyms')

    df_wkt = df_wkt.set_index(keys='Reference #')
    acronyms = acronyms.set_index(keys='Acronym')

    df_wkt = df_wkt.astype(str)
    df_wkt = df_wkt.applymap(lambda s: "" if s == 'nan' else s)
    acronyms = acronyms.astype(str)
    acronyms = acronyms.applymap(lambda s: "" if s == 'nan' else s)

    logger.debug(df_wkt.head())

    #
    doc = docx.Document()
    # doc.styles.add_style('cell reference', WD_STYLE_TYPE.PARAGRAPH)
    doc = add_acronyms(doc, df_wkt, lookup=acronyms)
    # doc = add_defined_terms(doc, df_wkt)
    doc = add_phase_breakdown(doc, df_wkt)
    doc = add_tasks(doc, df_wkt)
    doc = add_theme_breakdown(doc, df_wkt)
    doc = add_role_breakdown(doc, df_wkt)
    # doc = add_comments(doc, df_wkt, extract_comments('playbook_wkt.xlsx'))
    doc.save('playbook.docx')

    # extract_comments('playbook_wkt.xlsx')
    new_info_only_sheet('playbook_new.xlsx', df_wkt)
