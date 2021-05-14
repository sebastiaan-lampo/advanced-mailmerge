from docx import Document
# import docx
import docx.oxml
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)


def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True


def add_field(paragraph, field_code):
    # https://github.com/python-openxml/python-docx/issues/359#issuecomment-369271235
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def add_bookmark_pageref(paragraph, entry):
    field_code = f' PAGEREF "{entry}" \\h '
    add_field(paragraph, field_code)


def add_bookmark_ref(paragraph, entry):
    field_code = f' REF "{entry}" \\h '
    add_field(paragraph, field_code)


def set_cell_color(cell, color):
    # https://github.com/python-openxml/python-docx/issues/55
    tc = cell._tc
    tcPr = tc.tcPr
    color_xml = OxmlElement('w:shd')
    color_xml.set(qn('w:val'), 'clear')
    # color_xml.set(qn('w:color'), text)
    color_xml.set(qn('w:fill'), color)
    tcPr.append(color_xml)


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row
